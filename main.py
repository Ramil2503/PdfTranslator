import re
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, RGBColor
from deep_translator import GoogleTranslator
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
import os

def pdf_to_word(pdf_path, docx_path):
    """Convert PDF to Word, preserving font size and color."""
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()

def split_text(text, max_length=10000):
    """Split text into chunks without breaking sentences or words."""
    chunks = []
    sentences = re.split(r'(?<=[.!?]) +', text)  # Split by sentence boundaries

    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 <= max_length:
            if current_chunk:
                current_chunk += " "
            current_chunk += sentence
        else:
            chunks.append(current_chunk)
            current_chunk = sentence
    if current_chunk:
        chunks.append(current_chunk)

    return chunks

def translate_chunk(chunk, translator):
    """Translate a chunk of text."""
    try:
        return translator.translate(chunk)
    except Exception as e:
        print(f"Translation failed: {e}")
        return chunk  # Return the original chunk in case of an error

def translate_text(text, dest_language='ru', chunk_size=10000):
    """Translate text using GoogleTranslator, with optimal chunk size."""
    translator = GoogleTranslator(target=dest_language)
    chunks = split_text(text, max_length=chunk_size)
    translated_chunks = []

    # Parallel translation of text chunks
    with ThreadPoolExecutor(max_workers=10) as executor:
        future_to_chunk = {executor.submit(translate_chunk, chunk, translator): chunk for chunk in chunks}
        for future in tqdm(as_completed(future_to_chunk), desc="Translating text", total=len(chunks), unit="chunk"):
            result = future.result()
            if result is not None:
                translated_chunks.append(result)

    return ''.join(translated_chunks)

def rgb_color_to_tuple(rgb_color):
    """Convert RGBColor object to RGB tuple."""
    if rgb_color is None:
        return (0, 0, 0)  # Default to black if no color is set
    return (rgb_color[0], rgb_color[1], rgb_color[2])

def apply_style(run, font_size, font_color, bold, italic):
    """Apply font size, color, and styles to a run."""
    run.font.size = Pt(font_size)  # Set font size
    run.font.color.rgb = RGBColor(*font_color)  # Set font color
    run.bold = bold  # Set bold
    run.italic = italic  # Set italic

def translate_word(docx_path):
    """Translate text in the Word document and preserve font size, color, and style."""
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text = ""
            styles = []

            # Collect all runs and their styles
            for run in paragraph.runs:
                full_text += run.text
                styles.append({
                    'font_size': run.font.size.pt if run.font.size else 12,
                    'font_color': rgb_color_to_tuple(run.font.color.rgb) if run.font.color.rgb else (0, 0, 0),
                    'bold': run.bold if run.bold else False,
                    'italic': run.italic if run.italic else False
                })

            # Translate entire paragraph text at once
            translated_text = translate_text(full_text)
            
            if translated_text is None:
                translated_text = ""  # Ensure translated_text is a valid string

            # Clear existing runs
            paragraph.clear()

            # Create new runs with translated text
            translated_chunks = split_text(translated_text)
            for i, chunk in enumerate(translated_chunks):
                # Determine the style for this chunk
                style = styles[min(i, len(styles) - 1)]  # Ensure style index is within bounds
                new_run = paragraph.add_run(chunk)
                apply_style(new_run, style['font_size'], style['font_color'], style['bold'], style['italic'])

    doc.save(docx_path)

def word_to_pdf(docx_path, pdf_path):
    """Convert Word document to PDF."""
    from docx2pdf import convert as docx2pdf_convert
    try:
        docx2pdf_convert(docx_path, pdf_path)
    except Exception as e:
        print(f"Error converting Word to PDF: {e}")

def main():
    pdf_path = 'example_small.pdf'
    temp_docx_path = 'temp.docx'
    output_pdf_path = 'translated.pdf'

    # Step 1: Convert PDF to Word
    print("Converting PDF to Word...")
    pdf_to_word(pdf_path, temp_docx_path)
    
    # Step 2: Translate text in the Word document
    print("Translating text in Word document...")
    translate_word(temp_docx_path)
    
    # Step 3: Convert the modified Word document back to PDF
    print("Converting Word to PDF...")
    word_to_pdf(temp_docx_path, output_pdf_path)
    
    # Clean up temporary files
    os.remove(temp_docx_path)
    
    print(f'Translation completed and saved to {output_pdf_path}')

if __name__ == "__main__":
    main()